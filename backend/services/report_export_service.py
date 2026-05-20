# =============================================================================
# File: backend/services/report_export_service.py
# Project: Grocery Intelligence Platform
# Author: Anita Woodford
# Description: Generates PDF and DOCX report downloads from P1-19B report JSON.
# Security Note: Uses in-memory BytesIO exports; does not query MongoDB directly.
# SRS Traceability: SRS Section 16 DR-001 through DR-003, DR-020, DR-021.
# SDD Traceability: SDD Section 2.1; SDD Section 8; SDD Section 9; SDD Section 15.
# =============================================================================

from __future__ import annotations  # Enables modern Python type-hint behavior.

from io import BytesIO  # Provides in-memory file buffers.

from docx import Document  # Creates Word .docx documents.
from docx.shared import Inches  # Controls Word document layout spacing.


class ReportExportService:  # Defines the report export service.
    """Creates downloadable report files from already-generated report JSON."""  # Documents the service purpose.

    def build_pdf(self, report: dict) -> bytes:  # Builds PDF bytes from report data.
        from weasyprint import HTML  # Imports WeasyPrint lazily so app imports do not fail during test collection.

        html = self._build_html(report=report)  # Converts report JSON into printable HTML.
        pdf_bytes = HTML(string=html).write_pdf()  # Converts printable HTML into PDF bytes.
        return bytes(pdf_bytes)  # Returns immutable PDF bytes.

    def build_docx(self, report: dict) -> bytes:  # Builds DOCX bytes from report data.
        document = Document()  # Creates a new Word document.
        section = document.sections[0]  # Reads the default document section.
        section.top_margin = Inches(0.5)  # Sets readable top margin.
        section.bottom_margin = Inches(0.5)  # Sets readable bottom margin.
        section.left_margin = Inches(0.5)  # Sets readable left margin.
        section.right_margin = Inches(0.5)  # Sets readable right margin.

        document.add_heading("Weekly Deal Report", level=1)  # Adds the report title.
        document.add_paragraph(f"Week of: {report.get('week_of', 'Unknown')}")  # Adds the report week.
        document.add_paragraph(f"Generated at: {report.get('generated_at', 'Unknown')}")  # Adds the generation time.

        summary = report.get("summary", {})  # Reads report summary safely.
        document.add_heading("Weekly Summary", level=2)  # Adds the weekly summary heading.
        document.add_paragraph(f"Out-of-pocket: {self._money(summary.get('total_oop'))}")  # Shows OOP separately.
        document.add_paragraph(f"Register savings: {self._money(summary.get('total_register_savings'))}")  # Shows register savings.
        document.add_paragraph(f"Rewards earned: {self._money(summary.get('total_rewards_value'))}")  # Shows rewards as earnings.
        document.add_paragraph(f"Rebates back: {self._money(summary.get('total_rebates_back'))}")  # Shows rebates separately.
        document.add_paragraph(f"Final after rewards: {self._money(summary.get('final_after_rewards'))}")  # Shows final value as informational.
        document.add_paragraph("Note: Final after rewards is informational and is not the main register cost.")  # Clarifies final value.
        document.add_paragraph(f"Money-maker count: {summary.get('money_maker_count', 0)}")  # Shows money-maker count.

        for store in report.get("stores", []):  # Iterates through store sections.
            self._add_store_section(document=document, store=store)  # Adds one store section.

        buffer = BytesIO()  # Creates an in-memory output buffer.
        document.save(buffer)  # Writes the Word document into memory.
        return buffer.getvalue()  # Returns DOCX bytes.

    def _add_store_section(self, document: Document, store: dict) -> None:  # Adds one store section to DOCX.
        document.add_heading(f"Store: {store.get('store_ref', 'Unknown')}", level=2)  # Adds the store heading.

        totals = store.get("totals", {})  # Reads store totals safely.
        document.add_paragraph(f"Store out-of-pocket: {self._money(totals.get('total_oop'))}")  # Shows store OOP.
        document.add_paragraph(f"Store rewards earned: {self._money(totals.get('total_rewards_value'))}")  # Shows store rewards.
        document.add_paragraph(f"Store rebates back: {self._money(totals.get('total_rebates_back'))}")  # Shows store rebates.
        document.add_paragraph(f"Store final after rewards: {self._money(totals.get('final_after_rewards'))}")  # Shows display-only final.

        for deal in store.get("deals", []):  # Iterates through deal cards.
            self._add_deal_section(document=document, deal=deal)  # Adds one deal section.

    def _add_deal_section(self, document: Document, deal: dict) -> None:  # Adds one deal section to DOCX.
        title = f"Product: {deal.get('product_ref', 'Unknown')}"  # Builds product title text.
        if deal.get("is_money_maker"):  # Checks money-maker flag.
            title = f"{title} — MONEY MAKER"  # Adds money-maker badge text.

        document.add_heading(title, level=3)  # Adds the deal heading.
        document.add_paragraph(f"Register price: {self._money(deal.get('register_price'))}")  # Shows register price.
        document.add_paragraph(f"H-E-B coupons: {', '.join(deal.get('coupon_refs', [])) or 'None'}")  # Shows coupons separately.
        document.add_paragraph(f"Out-of-pocket: {self._money(deal.get('out_of_pocket'))}")  # Shows checkout cost.
        document.add_paragraph(f"Reward offers: {self._format_reward_offers(deal.get('reward_offers', []))}")  # Shows reward offers.
        document.add_paragraph(f"Receipt-required rewards: {self._format_required_rewards(deal.get('reward_offers', []), 'receipt_required')}")  # Shows receipt-required rewards.
        document.add_paragraph(f"Claim-required rewards: {self._format_required_rewards(deal.get('reward_offers', []), 'claim_required')}")  # Shows claim-required rewards.
        document.add_paragraph(f"Rebates back: {self._money(deal.get('total_rebates_back'))}")  # Shows rebates separately.
        document.add_paragraph(f"Final after rewards: {self._money(deal.get('final_after_rewards'))}")  # Shows informational final value.

    def _build_html(self, report: dict) -> str:  # Builds printable HTML for PDF output.
        summary = report.get("summary", {})  # Reads summary safely.
        store_sections = "".join(self._build_store_html(store=store) for store in report.get("stores", []))  # Builds all stores.

        return f"""<!doctype html>
<html>
<head>
<meta charset="utf-8">
<title>Weekly Deal Report</title>
<style>
body {{ font-family: Arial, sans-serif; font-size: 12px; color: #222; }}
h1 {{ font-size: 22px; margin-bottom: 4px; }}
h2 {{ font-size: 18px; border-bottom: 1px solid #ccc; padding-bottom: 4px; }}
h3 {{ font-size: 14px; margin-bottom: 4px; }}
.summary, .deal {{ border: 1px solid #ddd; padding: 10px; margin-bottom: 12px; }}
.grid {{ display: grid; grid-template-columns: repeat(3, 1fr); gap: 6px; }}
.label {{ font-weight: bold; }}
.badge {{ display: inline-block; padding: 2px 6px; border: 1px solid #444; font-weight: bold; }}
.note {{ font-size: 11px; color: #555; }}
</style>
</head>
<body>
<h1>Weekly Deal Report</h1>
<p>Week of: {self._escape(report.get("week_of", "Unknown"))}</p>
<p>Generated at: {self._escape(report.get("generated_at", "Unknown"))}</p>
<div class="summary">
<h2>Weekly Summary</h2>
<div class="grid">
<p><span class="label">Out-of-pocket:</span> {self._money(summary.get("total_oop"))}</p>
<p><span class="label">Register savings:</span> {self._money(summary.get("total_register_savings"))}</p>
<p><span class="label">Rewards earned:</span> {self._money(summary.get("total_rewards_value"))}</p>
<p><span class="label">Rebates back:</span> {self._money(summary.get("total_rebates_back"))}</p>
<p><span class="label">Final after rewards:</span> {self._money(summary.get("final_after_rewards"))}</p>
<p><span class="label">Money makers:</span> {summary.get("money_maker_count", 0)}</p>
</div>
<p class="note">Final after rewards is informational and is not the main register cost.</p>
</div>
{store_sections}
</body>
</html>"""  # Returns complete printable HTML.

    def _build_store_html(self, store: dict) -> str:  # Builds one store HTML section.
        totals = store.get("totals", {})  # Reads store totals safely.
        deals = "".join(self._build_deal_html(deal=deal) for deal in store.get("deals", []))  # Builds all deal cards.

        return f"""
<h2>Store: {self._escape(store.get("store_ref", "Unknown"))}</h2>
<div class="summary">
<p><span class="label">Store out-of-pocket:</span> {self._money(totals.get("total_oop"))}</p>
<p><span class="label">Store rewards earned:</span> {self._money(totals.get("total_rewards_value"))}</p>
<p><span class="label">Store rebates back:</span> {self._money(totals.get("total_rebates_back"))}</p>
<p><span class="label">Store final after rewards:</span> {self._money(totals.get("final_after_rewards"))}</p>
</div>
{deals}
"""  # Returns one store HTML section.

    def _build_deal_html(self, deal: dict) -> str:  # Builds one deal HTML block.
        badge = '<span class="badge">MONEY MAKER</span>' if deal.get("is_money_maker") else ""  # Builds money-maker badge.
        coupons = ", ".join(deal.get("coupon_refs", [])) or "None"  # Formats coupons separately.
        reward_offers = self._format_reward_offers(deal.get("reward_offers", []))  # Formats reward offers.
        receipt_rewards = self._format_required_rewards(deal.get("reward_offers", []), "receipt_required")  # Formats receipt-required rewards.
        claim_rewards = self._format_required_rewards(deal.get("reward_offers", []), "claim_required")  # Formats claim-required rewards.

        return f"""
<div class="deal">
<h3>Product: {self._escape(deal.get("product_ref", "Unknown"))} {badge}</h3>
<p><span class="label">Register price:</span> {self._money(deal.get("register_price"))}</p>
<p><span class="label">H-E-B coupons:</span> {self._escape(coupons)}</p>
<p><span class="label">Out-of-pocket:</span> {self._money(deal.get("out_of_pocket"))}</p>
<p><span class="label">Reward offers:</span> {self._escape(reward_offers)}</p>
<p><span class="label">Receipt-required rewards:</span> {self._escape(receipt_rewards)}</p>
<p><span class="label">Claim-required rewards:</span> {self._escape(claim_rewards)}</p>
<p><span class="label">Rebates back:</span> {self._money(deal.get("total_rebates_back"))}</p>
<p><span class="label">Final after rewards:</span> {self._money(deal.get("final_after_rewards"))}</p>
<p class="note">Rewards are earnings, not register deductions.</p>
</div>
"""  # Returns one deal HTML block.

    def _format_reward_offers(self, offers: list[dict]) -> str:  # Formats all reward offers.
        if not offers:  # Handles empty reward offers.
            return "None"  # Returns fallback display text.
        return "; ".join(str(offer.get("reward_provider", "Reward")) for offer in offers)  # Returns provider names.

    def _format_required_rewards(self, offers: list[dict], field_name: str) -> str:  # Formats required reward subsets.
        matching = [offer for offer in offers if bool(offer.get(field_name, False))]  # Filters matching offers.
        if not matching:  # Handles no matching rewards.
            return "None"  # Returns fallback display text.
        return "; ".join(str(offer.get("reward_provider", "Reward")) for offer in matching)  # Returns provider names.

    def _money(self, value: object) -> str:  # Formats numeric values as money.
        try:  # Starts safe conversion.
            return f"${float(value or 0):.2f}"  # Returns two-decimal currency text.
        except (TypeError, ValueError):  # Handles invalid numeric input.
            return "$0.00"  # Returns safe money fallback.

    def _escape(self, value: object) -> str:  # Escapes simple HTML-sensitive characters.
        text = str(value)  # Converts value to string.
        text = text.replace("&", "&amp;")  # Escapes ampersands.
        text = text.replace("<", "&lt;")  # Escapes less-than symbols.
        text = text.replace(">", "&gt;")  # Escapes greater-than symbols.
        text = text.replace('"', "&quot;")  # Escapes double quotes.
        return text.replace("'", "&#x27;")  # Escapes single quotes.